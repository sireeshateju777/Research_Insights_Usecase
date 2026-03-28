import docx
import sys

def read_docx(file_path):
    try:
        doc = docx.Document(file_path)
        content = []
        for para in doc.paragraphs:
            content.append(para.text)
        
        for table in doc.tables:
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text)
                content.append(" | ".join(row_data))
                
        with open('extracted_doc.txt', 'w', encoding='utf-8') as f:
            f.write('\n'.join(content))
            
        print("Success")
    except Exception as e:
        print(f"Error: {e}")

if __name__ == "__main__":
    read_docx(r'c:\Users\sireesha.malla\Videos\payslips\Research_Insights_AI_Platform.docx')
