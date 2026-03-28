from fastapi import FastAPI, UploadFile, File, BackgroundTasks, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import List, Optional
import os
import uuid
import datetime
import json
import re
import math
from collections import Counter

from dotenv import load_dotenv

load_dotenv()

# --- OpenAI setup ---
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "")
USE_AI = True  # Set to True once your OpenAI API key is configured

ai_client = None
if OPENAI_API_KEY:
    try:
        from openai import OpenAI
        print(f"[Init] OpenAI Client initialized with key: {OPENAI_API_KEY[:7]}...{OPENAI_API_KEY[-4:]}")
        ai_client = OpenAI(api_key=OPENAI_API_KEY)
    except Exception as e:
        print(f"[Warning] Could not initialize AI client: {e}")
else:
    print("[Warning] No OPENAI_API_KEY found. AI features will be disabled.")

app = FastAPI(title="Research Insights AI API")

# CORS for frontend
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Setup upload directory
UPLOAD_DIR = os.path.abspath(os.path.join(os.getcwd(), "..", "uploads"))
os.makedirs(UPLOAD_DIR, exist_ok=True)

# --- In-Memory Storage ---
documents_store = {}
insights_store = []


# --- Pydantic Models ---
class QueryRequest(BaseModel):
    document_ids: List[str]
    research_question: str
    response_format: str
    additional_context: Optional[str] = None

class InsightResponse(BaseModel):
    insight_id: str
    status: str
    format: str
    data: dict


# --- Text Extraction ---
def extract_text(file_path: str, filename: str) -> str:
    text = ""
    lower = filename.lower()
    try:
        if lower.endswith('.txt'):
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
        elif lower.endswith('.pdf'):
            import pypdf
            reader = pypdf.PdfReader(file_path)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        elif lower.endswith('.docx'):
            import docx
            doc = docx.Document(file_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    text += " | ".join(row_data) + "\n"
    except Exception as e:
        print(f"Error extracting text from {filename}: {e}")
    return text


# --- Chunking ---
def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        chunks.append(text[start:end])
        start += chunk_size - overlap
    return chunks


# --- Vector Similarity (Cosine Similarity) ---
def cosine_similarity(v1, v2):
    if not v1 or not v2: return 0.0
    dot_product = sum(a * b for a, b in zip(v1, v2))
    magnitude1 = math.sqrt(sum(a * a for a in v1))
    magnitude2 = math.sqrt(sum(a * a for a in v2))
    if magnitude1 == 0 or magnitude2 == 0:
        return 0.0
    return dot_product / (magnitude1 * magnitude2)


# --- Document Processing Pipeline ---
def process_document(file_id: str, file_path: str, filename: str):
    print(f"[Pipeline] Starting processing for: {filename} ({file_id})")
    documents_store[file_id]["status"] = "processing"

    text = extract_text(file_path, filename)
    if not text.strip():
        documents_store[file_id]["status"] = "failed"
        print(f"[Pipeline] No text extracted from {filename}")
        return

    chunks = chunk_text(text)
    print(f"[Pipeline] Created {len(chunks)} chunks from {filename}")

    # Generate Embeddings if AI is enabled
    embeddings = []
    if USE_AI and ai_client:
        try:
            print(f"[Pipeline] Generating embeddings for {len(chunks)} chunks via OpenAI...")
            # Using openai batch embeddings call
            response = ai_client.embeddings.create(
                model="text-embedding-3-small",
                input=chunks
            )
            embeddings = [item.embedding for item in response.data]
            print(f"[Pipeline] Embeddings generation successful.")
        except Exception as e:
            print(f"[Pipeline] Embedding Error: {e}")
            # Fallback to empty embeddings (will use keyword fallback)
            embeddings = [None] * len(chunks)

    documents_store[file_id]["chunks"] = chunks
    documents_store[file_id]["embeddings"] = embeddings
    documents_store[file_id]["full_text"] = text
    documents_store[file_id]["status"] = "completed"
    print(f"[Pipeline] Document {filename} processed successfully!")


# --- Retrieve relevant chunks ---
def get_query_embedding(query: str):
    if not USE_AI or not ai_client:
        return None
    try:
        response = ai_client.embeddings.create(
            model="text-embedding-3-small",
            input=query
        )
        return response.data[0].embedding
    except Exception as e:
        print(f"Error getting query embedding: {e}")
        return None


def retrieve_relevant_chunks(query: str, document_ids: List[str], top_k: int = 5) -> List[dict]:
    query_embedding = get_query_embedding(query)
    
    scored_chunks = []
    for doc_id in document_ids:
        doc = documents_store.get(doc_id)
        if not doc or doc["status"] != "completed":
            continue
        
        chunks = doc.get("chunks", [])
        embeddings = doc.get("embeddings", [])
        
        for i, chunk in enumerate(chunks):
            # Try vector similarity if embeddings exist
            score = 0
            if query_embedding and i < len(embeddings) and embeddings[i]:
                score = cosine_similarity(query_embedding, embeddings[i])
            else:
                # Keyword fallback score
                query_words = set(re.findall(r'\w+', query.lower()))
                chunk_words = set(re.findall(r'\w+', chunk.lower()))
                if query_words:
                    score = len(query_words & chunk_words) / len(query_words)
            
            scored_chunks.append({
                "text": chunk,
                "source": doc["filename"],
                "doc_id": doc_id,
                "chunk_index": i,
                "score": score
            })
            
    scored_chunks.sort(key=lambda x: x["score"], reverse=True)
    return scored_chunks[:top_k]


# --- Rule-Based Insight Generator (Fallback) ---
def clean_sentence(text: str) -> str:
    """Clean markdown formatting, bullets, and extra whitespace from a sentence."""
    text = re.sub(r'^#+\s*', '', text)           # Remove markdown headings
    text = re.sub(r'^\s*[-*•]\s*', '', text)     # Remove bullet points
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)  # Remove bold markdown
    text = re.sub(r'`([^`]+)`', r'\1', text)     # Remove code backticks
    text = re.sub(r'\[([^\]]+)\]', r'\1', text)  # Remove markdown links
    text = re.sub(r'\s+', ' ', text)             # Collapse whitespace
    return text.strip()


def is_meaningful_sentence(text: str) -> bool:
    """Filter out code fragments, schema definitions, and non-prose content."""
    skip_patterns = [
        r'^\s*-\s*`',           # Markdown code list items like "- `id`: UUID"
        r'UUID',                # Schema UUID references
        r'Primary Key',         # Schema definitions
        r'Foreign Key',         # Schema definitions
        r'VARCHAR|INTEGER|ENUM|TIMESTAMP|VECTOR|JSONB|ARRAY',  # SQL types
        r'^\s*\d+\.\s*$',      # Just numbered items
        r'^#+\s*\d',           # Markdown headings that are just numbers
        r'^\s*```',            # Code blocks
        r'^\s*\|',             # Table rows
        r'^\s*-\s*$',         # Empty bullet points
    ]
    for pattern in skip_patterns:
        if re.search(pattern, text, re.IGNORECASE):
            return False

    # Must have enough actual words (not just symbols/short fragments)
    words = re.findall(r'[a-zA-Z]{3,}', text)
    if len(words) < 5:
        return False

    # Filter out truncated sentences (cut mid-word by chunk boundaries)
    if text and not text[-1] in '.!?:)"' and len(text) > 50:
        last_word = text.split()[-1] if text.split() else ''
        if len(last_word) <= 2 or last_word.endswith('`') or last_word.endswith("'"):
            return False

    return True


def generate_rule_based_insights(query: str, chunks: List[dict], format_type: str) -> dict:
    """Generate structured insights from document chunks using text analysis."""
    all_text = "\n".join([c["text"] for c in chunks])
    raw_sentences = re.split(r'(?<=[.!?])\s+|\n+', all_text)

    sentences = []
    for s in raw_sentences:
        cleaned = clean_sentence(s.strip())
        if len(cleaned) > 30 and is_meaningful_sentence(cleaned):
            sentences.append(cleaned)

    seen = set()
    unique_sentences = []
    for s in sentences:
        s_key = s.lower()[:60]
        if s_key not in seen:
            seen.add(s_key)
            unique_sentences.append(s)
    sentences = unique_sentences

    finding_keywords = ['architecture', 'system', 'design', 'model', 'framework', 'built',
                        'serves', 'provides', 'enables', 'supports', 'consists', 'includes']
    want_keywords = ['must', 'should', 'shall', 'require', 'need', 'capability',
                     'scalability', 'latency', 'handle', 'concurrent']
    problem_keywords = ['risk', 'challenge', 'issue', 'concern', 'threat', 'prevent',
                        'failure', 'crash', 'limitation', 'constraint']
    action_keywords = ['recommend', 'suggest', 'plan', 'strategy', 'deploy', 'migrate',
                       'upgrade', 'optimize', 'improve']
    quickwin_keywords = ['performance', 'efficiency', 'reduce', 'batch', 'cache', 'index']

    def relevance_score(sentence, keywords):
        s_lower = sentence.lower()
        return sum(1 for kw in keywords if kw in s_lower)

    key_findings = []
    what_users_want = []
    common_problems = []
    next_steps = []
    quick_wins = []

    for s in sentences:
        scores = {
            'finding': relevance_score(s, finding_keywords),
            'want': relevance_score(s, want_keywords),
            'problem': relevance_score(s, problem_keywords),
            'action': relevance_score(s, action_keywords),
            'quickwin': relevance_score(s, quickwin_keywords),
        }
        best = max(scores, key=scores.get)
        if scores[best] == 0: continue
        if best == 'problem': common_problems.append(s)
        elif best == 'want': what_users_want.append(s)
        elif best == 'action': next_steps.append(s)
        elif best == 'quickwin': quick_wins.append(s)
        elif best == 'finding': key_findings.append(s)

    if not key_findings and sentences: key_findings = sentences[:3]

    key_findings = key_findings[:5]
    what_users_want = what_users_want[:4]
    common_problems = common_problems[:4]
    next_steps = next_steps[:4]
    quick_wins = quick_wins[:3]

    summary_parts = (key_findings[:2] + what_users_want[:1]) if key_findings else sentences[:3]
    exec_summary = ". ".join(summary_parts) + "." if summary_parts else "Analysis based on provided docs."
    exec_summary = re.sub(r'\.{2,}', '.', exec_summary)
    exec_summary = re.sub(r'\. \.', '.', exec_summary)

    if format_type.lower().startswith("simplified"):
        return {"bullet_points": (key_findings + what_users_want + common_problems)[:8]}

    result = {}
    if key_findings: result["key_findings"] = [{"text": t, "priority": "high"} for t in key_findings]
    if what_users_want: result["what_users_want"] = [{"text": t, "priority": "high"} for t in what_users_want]
    if quick_wins: result["strategic_quick_wins"] = [{"text": t, "priority": "medium"} for t in quick_wins]
    if common_problems: result["common_problems"] = [{"text": t, "priority": "medium"} for t in common_problems]
    if next_steps: result["recommended_next_steps"] = [{"text": t, "priority": "high"} for t in next_steps]
    result["executive_summary"] = exec_summary
    return result


# --- Import Agent Pipeline ---
from agent import run_pipeline


# --- API Endpoints ---

@app.get("/documents")
async def list_documents():
    docs = []
    for doc_id, doc in documents_store.items():
        docs.append({
            "id": doc_id,
            "filename": doc["filename"],
            "status": doc["status"],
            "upload_date": doc["upload_date"],
            "chunk_count": len(doc.get("chunks", []))
        })
    return {"documents": docs}


@app.post("/upload")
async def upload_document(background_tasks: BackgroundTasks, file: UploadFile = File(...)):
    file_id = str(uuid.uuid4())
    file_path = os.path.join(UPLOAD_DIR, f"{file_id}_{file.filename}")

    with open(file_path, "wb") as buffer:
        buffer.write(await file.read())

    documents_store[file_id] = {
        "filename": file.filename,
        "file_path": file_path,
        "status": "pending",
        "upload_date": datetime.datetime.now().isoformat(),
        "chunks": [],
        "embeddings": [],
        "full_text": ""
    }

    background_tasks.add_task(process_document, file_id, file_path, file.filename)

    return {
        "id": file_id,
        "filename": file.filename,
        "status": "pending",
        "message": "File uploaded and queued for processing."
    }


@app.delete("/documents/{doc_id}")
async def delete_document(doc_id: str):
    if doc_id not in documents_store:
        raise HTTPException(status_code=404, detail="Document not found")
    doc = documents_store.pop(doc_id)
    if os.path.exists(doc["file_path"]):
        os.remove(doc["file_path"])
    return {"message": f"Document {doc['filename']} deleted."}


@app.post("/generate_insights", response_model=InsightResponse)
async def generate_insights(request: QueryRequest):
    if not request.document_ids or not request.research_question:
        raise HTTPException(status_code=400, detail="Missing required fields.")

    for doc_id in request.document_ids:
        doc = documents_store.get(doc_id)
        if not doc:
            raise HTTPException(status_code=404, detail=f"Document {doc_id} not found.")
        if doc["status"] != "completed":
            raise HTTPException(status_code=400, detail=f"Document {doc['filename']} is still {doc['status']}.")

    insight_id = str(uuid.uuid4())

    # Retrieve relevant chunks
    relevant_chunks = retrieve_relevant_chunks(request.research_question, request.document_ids)

    # Generate insights
    if USE_AI and ai_client:
        try:
            # Use LangGraph + OpenAI pipeline (GPT-4o)
            structured_data = run_pipeline(
                query=request.research_question,
                document_ids=request.document_ids,
                format_type=request.response_format,
                retrieved_chunks=relevant_chunks,
                client=ai_client
            )
        except Exception as e:
            print(f"AI Generation Error: {e}")
            structured_data = generate_rule_based_insights(request.research_question, relevant_chunks, request.response_format)
            structured_data["executive_summary"] = f"(AI Fallback) {structured_data['executive_summary']} (Reason: {str(e)})"
    else:
        # Use rule-based text analysis (no API key needed)
        structured_data = generate_rule_based_insights(
            request.research_question,
            relevant_chunks,
            request.response_format
        )

    # Add citations
    seen_sources = set()
    citations = []
    for chunk in relevant_chunks:
        source = chunk.get("source", "Unknown")
        if source not in seen_sources:
            seen_sources.add(source)
            citations.append({"id": len(citations) + 1, "source": source, "doc_id": chunk.get("doc_id", "")})
    structured_data["sources_and_citations"] = citations

    # Save to history
    insights_store.append({"id": insight_id, "query": request.research_question, "date": datetime.date.today().isoformat(), "docs_included": len(request.document_ids), "format": request.response_format, "data": structured_data})

    return InsightResponse(insight_id=insight_id, status="completed", format=request.response_format, data=structured_data)


@app.get("/export/{insight_id}")
async def export_insight(insight_id: str, format_type: str = "docx"):
    insight = next((item for item in insights_store if item["id"] == insight_id), None)
    if not insight: raise HTTPException(status_code=404, detail="Insight not found")
    data = insight["data"]

    if format_type.lower() == "docx":
        import docx
        from fastapi.responses import FileResponse
        doc = docx.Document()
        doc.add_heading("Research Insights Report", 0)
        doc.add_paragraph(f"Question: {insight['query']}")
        doc.add_paragraph(f"Date: {insight['date']}")
        doc.add_paragraph("")

        section_map = {"key_findings": "Key Findings", "what_users_want": "What Users Want", "strategic_quick_wins": "Strategic Quick Wins", "common_problems": "Common Problems", "recommended_next_steps": "Recommended Next Steps"}
        for key, title in section_map.items():
            items = data.get(key, [])
            if items:
                doc.add_heading(title, level=1)
                for item in items:
                    text = item["text"] if isinstance(item, dict) else item
                    priority = item.get("priority", "") if isinstance(item, dict) else ""
                    doc.add_paragraph(f"• {text} [{priority}]", style="List Bullet")
        
        if data.get("executive_summary"):
            doc.add_heading("Executive Summary", level=1)
            doc.add_paragraph(data["executive_summary"])
        
        if data.get("sources_and_citations"):
            doc.add_heading("Sources & Citations", level=1)
            for c in data["sources_and_citations"]: doc.add_paragraph(f"[{c['id']}] {c['source']}")

        export_path = os.path.join(UPLOAD_DIR, f"insight_{insight_id}.docx")
        doc.save(export_path)
        return FileResponse(export_path, filename=f"Research_Insights_{insight['date']}.docx", media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    return HTTPException(status_code=400, detail="Only DOCX export supported.")


@app.get("/history")
async def get_insight_history(date_filter: Optional[str] = None, format_filter: Optional[str] = None):
    results = insights_store
    if date_filter: results = [r for r in results if r["date"] == date_filter]
    if format_filter: results = [r for r in results if r["format"].lower() == format_filter.lower()]
    return {"items": results, "total": len(results), "page": 1}


if __name__ == "__main__":
    import uvicorn
    print("\n=== Research Insights AI Platform ===")
    print(f"AI Mode: Enabled (OpenAI)")
    print(f"LLM: GPT-4o | Embeddings: text-embedding-3-small")
    print(f"Server: http://localhost:8001\n")
    uvicorn.run(app, host="0.0.0.0", port=8001)

